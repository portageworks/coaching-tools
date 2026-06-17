"""
Converts the structured resume text output from the AI into a formatted .docx.

Resume format expected:
  Line 1: Candidate Full Name
  Line 2: Contact info
  ## SECTION HEADER
  ### Company Name (in EXPERIENCE state)
  **Title** | Date range  (title line -- contains a year)
  SCOPE: ...  (italic scope paragraph)
  - Bullet
  Body paragraphs (SUMMARY, ADDITIONAL, EDUCATION, CERTIFICATIONS)
"""
import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── colour palette ────────────────────────────────────────────────────────────
_CHARCOAL  = RGBColor(0x1e, 0x20, 0x22)
_SLATE     = RGBColor(0x3a, 0x5a, 0x7c)
_TEXT_MID  = RGBColor(0x4a, 0x50, 0x58)
_GREY_FILL = "F2F2F2"
_SLATE_HEX = "3A5A7C"

DEFAULT_FONT = "Calibri"
ALLOWED_FONTS = {"Arial", "Calibri", "Cambria", "Times New Roman"}


def _run(paragraph, text, bold=False, italic=False, size_pt=11,
         color=None, highlight=False, font=DEFAULT_FONT):
    run = paragraph.add_run(text)
    run.font.name  = font
    run.font.size  = Pt(size_pt)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    if highlight:
        rPr = run._r.get_or_add_rPr()
        hl  = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "yellow")
        rPr.append(hl)
    return run


def _add_runs_with_placeholders(paragraph, text, bold=False, italic=False,
                                 size_pt=11, color=None, font=DEFAULT_FONT):
    """Split on ((placeholder)) markers and highlight them yellow."""
    parts = re.split(r"(\(\([^)]*?\)\))", text)
    for part in parts:
        if not part:
            continue
        is_ph = part.startswith("((") and part.endswith("))")
        _run(paragraph, part, bold=bold, italic=italic, size_pt=size_pt,
             color=color, highlight=is_ph, font=font)


def _strip_bold(text):
    return re.sub(r"\*\*([^*]*)\*\*", r"\1", text)


def _set_para_border_bottom(paragraph, color="AAAAAA", size=4, space=1):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), str(space))
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_shading(paragraph, fill=_GREY_FILL):
    pPr  = paragraph._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)


def _is_title_line(line):
    return bool(re.search(r"\b(19|20)\d{2}\b", line))


def _split_title_date(line):
    stripped = _strip_bold(line).strip()
    pipe = stripped.rfind(" | ")
    if pipe >= 0:
        return stripped[:pipe].strip(), stripped[pipe + 3:].strip()
    m = re.match(r"^(.+?)\s+((?:19|20)\d{2}.{0,25})$", stripped)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return stripped, ""


# ── paragraph builders ────────────────────────────────────────────────────────

def _name_para(doc, text, font=DEFAULT_FONT):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    _run(p, _strip_bold(text).upper(), bold=True, size_pt=20,
         color=_CHARCOAL, font=font)
    return p


def _contact_para(doc, text, font=DEFAULT_FONT):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    _add_runs_with_placeholders(p, _strip_bold(text), size_pt=10,
                                 color=_TEXT_MID, font=font)
    return p


def _section_header(doc, text, font=DEFAULT_FONT):
    label = re.sub(r"^#+\s*", "", text).strip().upper()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    _set_shading(p, _GREY_FILL)
    _set_para_border_bottom(p, color=_SLATE_HEX, size=6)
    _run(p, label, bold=True, size_pt=11, color=_CHARCOAL, font=font)
    return p


def _company_para(doc, text, font=DEFAULT_FONT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0)
    _add_runs_with_placeholders(p, _strip_bold(text), bold=True,
                                 size_pt=11, color=_CHARCOAL, font=font)
    return p


def _title_para(doc, line, font=DEFAULT_FONT):
    title, date = _split_title_date(line)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),  "right")
    tab.set(qn("w:pos"),  "9360")
    tabs.append(tab)
    pPr.append(tabs)
    _add_runs_with_placeholders(p, title, bold=True, size_pt=11,
                                 color=_CHARCOAL, font=font)
    if date:
        _run(p, "\t", font=font, size_pt=11)
        _add_runs_with_placeholders(p, date, size_pt=11, color=_TEXT_MID, font=font)
    return p


def _scope_para(doc, text, font=DEFAULT_FONT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    _add_runs_with_placeholders(p, text, italic=True, size_pt=11,
                                 color=_TEXT_MID, font=font)
    return p


def _bullet_para(doc, text, font=DEFAULT_FONT):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.2)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(2)
    _add_runs_with_placeholders(p, text, size_pt=11, color=_TEXT_MID, font=font)
    return p


def _body_para(doc, text, centered=False, font=DEFAULT_FONT):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _add_runs_with_placeholders(p, _strip_bold(text), size_pt=11,
                                 color=_TEXT_MID, font=font)
    return p


# ── main builder ──────────────────────────────────────────────────────────────

def resume_to_docx(resume_text: str, font: str = DEFAULT_FONT) -> bytes:
    if font not in ALLOWED_FONTS:
        font = DEFAULT_FONT

    doc = Document()

    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(11)

    lines  = resume_text.splitlines()
    state  = "HEADER"
    header_count    = 0
    in_role_block   = False
    after_title     = False
    seen_bullet     = False
    scope_buffer    = []

    def flush_scope():
        nonlocal scope_buffer
        if scope_buffer:
            _scope_para(doc, " ".join(scope_buffer).strip(), font=font)
            scope_buffer = []

    for line in lines:
        raw  = line
        line = line.strip()

        if line.startswith("## "):
            flush_scope()
            in_role_block = False
            after_title   = False
            seen_bullet   = False
            sec = re.sub(r"^##\s*", "", line).strip().upper()
            _section_header(doc, line, font=font)
            if "ADDITIONAL" in sec:
                state = "ADDITIONAL"
            elif "EXPERIENCE" in sec:
                state = "EXPERIENCE"
            elif "COMPETENC" in sec:
                state = "COMPETENCIES"
            elif "EDUCATION" in sec:
                state = "EDUCATION"
            elif "CERTIF" in sec or "PROFESSIONAL DEV" in sec:
                state = "CERTIFICATIONS"
            elif "LANGUAGE" in sec:
                state = "LANGUAGES"
            else:
                state = "SUMMARY"
            continue

        if line.startswith("### ") and state == "EXPERIENCE":
            flush_scope()
            in_role_block = True
            after_title   = False
            seen_bullet   = False
            _company_para(doc, line[4:].strip(), font=font)
            continue

        if not line:
            if seen_bullet:
                flush_scope()
            continue

        if re.match(r"^[-*]\s+", line) and state == "EXPERIENCE":
            if not seen_bullet:
                flush_scope()
                seen_bullet = True
            _bullet_para(doc, re.sub(r"^[-*]\s+", "", line), font=font)
            continue

        if state == "HEADER":
            clean = re.sub(r"^#+\s*", "", line)
            if header_count == 0:
                _name_para(doc, clean, font=font)
            else:
                _contact_para(doc, clean, font=font)
            header_count += 1
            if header_count >= 2:
                state = "PRE_SECTION"
            continue

        if state in ("SUMMARY", "PRE_SECTION"):
            _body_para(doc, line, font=font)
            continue

        if state == "COMPETENCIES":
            _body_para(doc, line, centered=True, font=font)
            continue

        if state == "EXPERIENCE":
            if _is_title_line(line):
                flush_scope()
                after_title = True
                seen_bullet = False
                if not in_role_block:
                    in_role_block = True
                _title_para(doc, line, font=font)
            elif not in_role_block:
                _company_para(doc, line, font=font)
                in_role_block = True
            elif after_title and not seen_bullet:
                scope_line = re.sub(r"^SCOPE:\s*", "", line, flags=re.IGNORECASE).strip()
                scope_buffer.append(scope_line)
            else:
                _body_para(doc, line, font=font)
            continue

        if state == "ADDITIONAL":
            _body_para(doc, line, font=font)
            continue

        _body_para(doc, line, font=font)

    flush_scope()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
