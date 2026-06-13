"""
Converts a markdown string into a .docx BytesIO object.
Handles: headings (# ## ###), bold (**text**), bullet lists (- item),
blockquotes (> text), horizontal rules (---), and paragraphs.
"""
import re
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip("#")
    return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def _add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "6")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), "3A5A7C")
    pBdr.append(left)
    pPr.append(pBdr)
    _add_runs(p, text, size_pt=11)


def _add_runs(paragraph, text, size_pt=11, bold_all=False, color_hex=None):
    """Split text on **bold** markers and add runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run_text = part[2:-2] if is_bold else part
        run = paragraph.add_run(run_text)
        run.font.name = "Arial"
        run.font.size = Pt(size_pt)
        run.bold = is_bold or bold_all
        if color_hex:
            r, g, b = _hex_to_rgb(color_hex)
            run.font.color.rgb = RGBColor(r, g, b)


def _set_heading_style(paragraph, level):
    configs = {
        1: (18, True, "1E2022", 16, 8),
        2: (14, True, "3A5A7C", 12, 6),
        3: (12, True, "1E2022", 10, 4),
    }
    size_pt, bold, color, space_before, space_after = configs.get(level, (11, False, "000000", 6, 3))
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size_pt)
        run.bold = bold
        r, g, b = _hex_to_rgb(color)
        run.font.color.rgb = RGBColor(r, g, b)


def markdown_to_docx(markdown_text: str) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = markdown_text.splitlines()
    i = 0
    in_blockquote = False
    bq_lines = []

    def flush_blockquote():
        nonlocal bq_lines
        if bq_lines:
            _add_blockquote(doc, " ".join(bq_lines))
            bq_lines = []

    while i < len(lines):
        line = lines[i]

        # Blockquote
        if line.startswith("> "):
            in_blockquote = True
            bq_lines.append(line[2:])
            i += 1
            continue
        elif line.startswith(">"):
            in_blockquote = True
            bq_lines.append(line[1:].strip())
            i += 1
            continue
        else:
            if in_blockquote:
                flush_blockquote()
                in_blockquote = False

        # Heading
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Strip italic markers for display
            text = re.sub(r"\*([^*]+)\*", r"\1", text)
            p = doc.add_paragraph()
            _add_runs(p, text, size_pt=[18, 14, 12][level - 1], bold_all=True,
                      color_hex=["1E2022", "3A5A7C", "1E2022"][level - 1])
            _set_heading_style(p, level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", line.strip()) or re.match(r"^_{3,}$", line.strip()):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Bullet list item
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            text = re.sub(r"^[-*]\s+", "", line)
            _add_runs(p, text, size_pt=11)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        # Collect consecutive non-empty, non-special lines
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") \
                and not lines[i].startswith(">") and not re.match(r"^[-*]\s+", lines[i]) \
                and not re.match(r"^-{3,}$", lines[i].strip()):
            para_lines.append(lines[i])
            i += 1

        text = " ".join(para_lines)
        # Strip italic wrappers (e.g. *Setup -- say this before asking:*)
        text = re.sub(r"(?<!\*)\*(?!\*)([^*]+)(?<!\*)\*(?!\*)", r"\1", text)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_runs(p, text, size_pt=11)

    flush_blockquote()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
