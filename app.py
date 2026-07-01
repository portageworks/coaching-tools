from gevent import monkey
monkey.patch_all()

import os
import re
import json
import io
import time
import base64
import queue
import threading
import base64
import hmac
import secrets as _secrets
import urllib.request
import urllib.parse
from collections import deque, defaultdict
from datetime import timedelta
from flask import (Flask, render_template, render_template_string, request, jsonify,
                   Response, stream_with_context, send_file, session, redirect,
                   url_for, abort)
import anthropic

from prompts.session_generator import SYNTHESIS_SYSTEM, COACH_GUIDE_SYSTEM, WORKSHEET_SYSTEM
from prompts.session_builder import (
    interview_program_prompt, stories_prompt,
    positioning_prompt, resume_diagnostic_prompt, resume_rewrite_prompt,
)
from prompts.resume_tool import DIAGNOSTIC_SYSTEM, REWRITE_SYSTEM
from prompts.session_builder_part2 import (
    summary_prompt, roleplay_a_prompt, roleplay_b_prompt,
    branding_prompt, training_prompt,
    linkedin_strategy_prompt, coach_handoff_prompt,
)
from docx_builder import markdown_to_docx
from resume_docx_builder import resume_to_docx
from pdf_builder import (
    build_client_html, build_positioning_html, build_strategy_package_html,
    build_worksheet_html, build_worksheet_cue_html, render_pdf,
    positioning_warnings, package_warnings,
)
import cue_store
import segno
import strategy_store

# Load .env for local dev so ANTHROPIC_API_KEY (etc.) is available when running
# `python app.py` directly. setdefault means real environment variables — as set
# in production (Railway) — always win, so this is a no-op there. Kept dependency-
# free so it works even where python-dotenv isn't installed.
_envfile = os.path.join(os.path.dirname(__file__), ".env")
if os.path.exists(_envfile):
    for _line in open(_envfile, encoding="utf-8"):
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            os.environ.setdefault(_k.strip(), _v.strip().strip('"').strip("'"))

app = Flask(__name__)

# ── Access control ──────────────────────────────────────────────────────────────
# A single shared password gates the whole site (it holds client PII). On success
# we set a signed, 30-day cookie so the coach logs in once per device. Set
# APP_PASSWORD in the environment to enable; if it's unset (e.g. local dev), the
# gate is off. SECRET_KEY signs the cookie — set it in Railway so the login
# survives redeploys (otherwise a fresh random key logs everyone out each deploy).
APP_PASSWORD = os.environ.get("APP_PASSWORD", "")
app.secret_key = os.environ.get("SECRET_KEY") or _secrets.token_hex(32)
app.permanent_session_lifetime = timedelta(days=30)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=bool(APP_PASSWORD),  # cookie only over HTTPS in prod
)

LOGIN_HTML = """<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Sign in</title><style>
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:-apple-system,Segoe UI,Roboto,sans-serif;background:#f5f3f0;color:#2c3035;
  display:flex;align-items:center;justify-content:center;min-height:100vh;padding:24px;}
.box{background:#fff;border:1px solid #e0dbd4;border-radius:8px;padding:36px 38px;width:100%;max-width:360px;}
h1{font-family:'Roboto Slab',serif;font-size:20px;color:#1e2022;margin-bottom:6px;}
p{font-size:13px;color:#7a8088;margin-bottom:20px;}
input{width:100%;font-size:15px;padding:12px 14px;border:1px solid #d8d3cc;border-radius:6px;margin-bottom:14px;}
input:focus{outline:none;border-color:#3a5a7c;}
button{width:100%;font-size:15px;font-weight:600;padding:12px;border:none;border-radius:6px;
  background:#1e2022;color:#fff;cursor:pointer;}
.err{font-size:13px;color:#8b4040;margin-bottom:12px;}
</style></head><body><div class="box">
<h1>Coaching Tools</h1><p>Enter the password to continue.</p>
{% if error %}<div class="err">{{ error }}</div>{% endif %}
<form method="POST" action="{{ url_for('login') }}">
  <input type="hidden" name="next" value="{{ next }}">
  <input type="password" name="password" placeholder="Password" autofocus autocomplete="current-password">
  <button type="submit">Sign in</button>
</form></div></body></html>"""


# The client-facing prompt library is served on its own public host
# (promptrunway.work). On THAT host only the library is reachable — the coaching
# tools 404 — and the library is open, with no coach login. Every other host is
# the private side, behind the shared-password gate below. Set PUBLIC_HOST in the
# environment if the client domain ever changes.
PUBLIC_HOST = os.environ.get("PUBLIC_HOST", "promptrunway.work").lower()
_PUBLIC_PREFIXES = ("/prompt-library", "/api/prompt-library", "/static/")

# Cloudflare Turnstile (bot protection on the generate endpoint). Set both env
# vars to require a token on every generation; leave them unset and verification
# is skipped, so the tool works unchanged until you configure Cloudflare.
TURNSTILE_SITE_KEY = os.environ.get("TURNSTILE_SITE_KEY", "")
TURNSTILE_SECRET = os.environ.get("TURNSTILE_SECRET", "")
_TURNSTILE_VERIFY_URL = "https://challenges.cloudflare.com/turnstile/v0/siteverify"


@app.before_request
def _gate():
    host = request.host.split(":")[0].lower()
    if host == PUBLIC_HOST:
        # Client domain: expose only the library (the bare root serves it too);
        # everything else is invisible.
        if request.path == "/" or request.path.startswith(_PUBLIC_PREFIXES):
            return
        abort(404)

    # Private host: single shared password gate.
    if not APP_PASSWORD:
        return  # gate disabled until a password is configured
    p = request.path
    if p == "/login" or p.startswith("/static/"):
        return
    if session.get("authed"):
        return
    if p.startswith("/api/"):
        abort(401)  # JSON callers get a clean 401, not an HTML redirect
    return redirect(url_for("login", next=p))


@app.route("/login", methods=["GET", "POST"])
def login():
    if not APP_PASSWORD:
        return redirect("/")
    error = ""
    nxt = request.values.get("next", "/")
    if not nxt.startswith("/"):
        nxt = "/"  # guard against open-redirect
    if request.method == "POST":
        if hmac.compare_digest(request.form.get("password", ""), APP_PASSWORD):
            session.permanent = True
            session["authed"] = True
            return redirect(nxt)
        error = "Incorrect password."
    return render_template_string(LOGIN_HTML, error=error, next=nxt)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.context_processor
def _inject_auth():
    # Lets templates show a "Sign out" link only when the password gate is active.
    return {"auth_enabled": bool(APP_PASSWORD)}

# max_retries handles transient API failures (overloaded, rate limits, timeouts,
# 5xx) with exponential backoff so a run succeeds the first time instead of
# erroring out and forcing a full re-run. Default is 2; raised for reliability.
client = anthropic.Anthropic(
    api_key=os.environ.get("ANTHROPIC_API_KEY"),
    max_retries=5,
)

# ── Model selection ───────────────────────────────────────────────────────────
# SMART = higher-quality, client-facing deliverables (resume rewrites, branding,
#         session guides, roleplay). FAST = cheap internal jobs (diagnostics,
#         coach-only training notes). Flip these in one place.
def msg_text(resp):
    """Safely pull the text out of a messages response. The API can occasionally
    return a message with no content blocks (or non-text blocks); indexing
    content[0] blindly throws 'list index out of range'. Join all text blocks
    and return '' if there are none."""
    return "".join(
        getattr(b, "text", "") for b in (resp.content or []) if getattr(b, "type", None) == "text"
    ).strip()


def complete_text(system, user_content, model, max_tokens, attempts=3):
    """Call the model and return its text, retrying if the reply comes back
    empty (an intermittent API condition that otherwise crashed jobs with
    'list index out of range'). Raises ValueError if still empty after
    `attempts` tries. user_content may be a string or a messages list."""
    messages = ([{"role": "user", "content": user_content}]
                if isinstance(user_content, str) else user_content)
    for _ in range(attempts):
        resp = client.messages.create(
            model=model, max_tokens=max_tokens, system=system, messages=messages,
        )
        text = msg_text(resp)
        if text:
            return text
    raise ValueError("Model returned an empty response after multiple attempts. "
                     "This is usually transient — please run it again.")


def complete_json(system, user_content, model, max_tokens, attempts=3):
    """Call the model and parse a JSON object from its reply. If the reply
    isn't valid JSON, hand the model its own output back with a correction and
    re-ask, up to `attempts` times. This catches malformed-JSON failures, which
    the client's transient-error retries do not. Raises ValueError if it never
    returns parseable JSON."""
    messages = [{"role": "user", "content": user_content}]
    last_raw = ""
    for _ in range(attempts):
        resp = client.messages.create(
            model=model, max_tokens=max_tokens, system=system, messages=messages,
        )
        last_raw = msg_text(resp)
        cleaned = last_raw.replace("```json", "").replace("```", "").strip()
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            messages = [
                {"role": "user", "content": user_content},
                {"role": "assistant", "content": last_raw},
                {"role": "user", "content": "That was not valid JSON. Reply with ONLY "
                 "the JSON object — no preface, no explanation, no markdown code fences."},
            ]
    raise ValueError(f"Model did not return valid JSON after {attempts} attempts. "
                     f"Last reply began: {last_raw[:200]}")


MODEL_SMART = "claude-opus-4-8"
MODEL_FAST  = "claude-sonnet-4-6"


# ── Generic endpoints ─────────────────────────────────────────────────────────

@app.route("/")
def index():
    # On the client domain the bare root IS the prompt library; the private host
    # keeps the coaching hub.
    if request.host.split(":")[0].lower() == PUBLIC_HOST:
        return _render_prompt_library()
    return render_template("index.html")


@app.route("/api/chat", methods=["POST"])
def chat():
    data = request.get_json()
    messages = data.get("messages", [])
    system = data.get("system", "")
    model = data.get("model", "claude-sonnet-4-6")
    max_tokens = data.get("max_tokens", 8096)

    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=system,
        messages=messages,
    )
    return jsonify({"content": msg_text(response)})


# ── Resume Tool ───────────────────────────────────────────────────────────────

@app.route("/resume-tool")
def resume_tool():
    return render_template("resume_tool.html")


@app.route("/api/resume/diagnose", methods=["POST"])
def resume_diagnose():
    data       = request.get_json()
    resume     = (data.get("resume") or "").strip()
    transcript = (data.get("transcript") or "").strip()
    target     = (data.get("target") or "").strip()

    if not resume:
        return jsonify({"error": "Resume text is required."}), 400

    user_msg = f"RESUME:\n\n{resume}"
    if target:
        user_msg += f"\n\n---\nTARGET POSITIONING:\n\n{target}"
    if transcript:
        user_msg += f"\n\n---\nSESSION TRANSCRIPT (Primary Truth — overrides resume where they conflict):\n\n{transcript}"

    try:
        result = complete_json(DIAGNOSTIC_SYSTEM, user_msg, MODEL_FAST, 4000)
    except Exception as e:
        return jsonify({"error": f"Evaluation failed: {e}"}), 500
    return jsonify(result)


@app.route("/api/resume/rewrite", methods=["POST"])
def resume_rewrite():
    data       = request.get_json()
    resume     = (data.get("resume") or "").strip()
    transcript = (data.get("transcript") or "").strip()
    target     = (data.get("target") or "").strip()
    diag       = data.get("diagnostic") or {}

    if not resume:
        return jsonify({"error": "Resume text is required."}), 400

    level = diag.get("level", "EXECUTIVE")
    score = diag.get("readinessScore", 5)
    mode  = diag.get("revisionMode", "REWRITE")

    user_msg  = f"CANDIDATE LEVEL: {level}\n"
    user_msg += f"READINESS SCORE: {score}/10\n"
    user_msg += f"REVISION MODE: {mode}\n\n"
    user_msg += "STRATEGY BRIEF:\n"
    user_msg += f"Strengths to Preserve: {diag.get('strengths', '')}\n"
    user_msg += f"Narrative Arc: {diag.get('narrativeArc', '')}\n"
    user_msg += f"Impact Opportunities: {diag.get('impactOpportunities', '')}\n"
    user_msg += f"Keyword Alignment: {diag.get('keywordAlignment', '')}\n"
    user_msg += f"Positioning Hypothesis: {diag.get('positioningHypothesis', '')}\n\n"
    user_msg += f"RESUME:\n\n{resume}"
    if target:
        user_msg += f"\n\n---\nTARGET POSITIONING:\n\n{target}"
    if transcript:
        user_msg += f"\n\n---\nSESSION TRANSCRIPT (Primary Truth):\n\n{transcript}"

    content = complete_text(REWRITE_SYSTEM, user_msg, MODEL_SMART, 8000)
    return jsonify({"content": content})


@app.route("/api/resume/download.docx", methods=["POST"])
def resume_download_docx():
    data        = request.get_json()
    resume_text = data.get("content", "")
    font        = (data.get("font") or "Calibri").strip()

    docx_bytes = resume_to_docx(resume_text, font=font)

    first_line = resume_text.strip().splitlines()[0] if resume_text.strip() else ""
    import re as _re
    name = _re.sub(r"[^a-zA-Z\s'-]", "", first_line).strip()
    parts = name.split()
    if len(parts) >= 2:
        filename = f"{parts[-1]}_{parts[0]}_Resume.docx"
    elif parts:
        filename = f"{parts[0]}_Resume.docx"
    else:
        filename = "Resume.docx"

    return send_file(
        io.BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ── Prompt Library (client-facing) ────────────────────────────────────────────
# Prompts live as data in prompt_library/prompts.json (extracted from the
# original static promptrunway.work page). The page renders from this list, so
# adding or editing a prompt is a data change, not markup surgery. Each prompt's
# "model" ("fast"=Sonnet / "smart"=Opus) drives routing when inline generation
# lands in Phase 3.
_PROMPT_LIB_LABELS = {
    "resume": "Resume", "cover": "Cover Letters", "search": "Job Search",
    "network": "Networking", "interview": "Interview", "salary": "Salary",
    "pivot": "Career Pivots", "mindset": "Mindset",
}


def _load_prompt_library():
    path = os.path.join(os.path.dirname(__file__), "prompt_library", "prompts.json")
    with open(path, encoding="utf-8") as f:
        prompts = json.load(f)
    cats, order = {}, []
    for p in prompts:
        t = p["template"]
        tl = t.lower()
        # Input flags are derived from the template, so an authored prompt only
        # needs its text — no by-hand bookkeeping to keep in sync.
        p["needs_resume"] = "attach your resume" in tl or "paste your resume" in tl
        p["needs_paste"] = "[paste" in tl
        p["needs_job_description"] = "[PASTE THE JOB DESCRIPTION HERE]" in t
        p["needs_cover_letter"] = "[PASTE YOUR COVER LETTER HERE]" in t
        # extra_inputs: optional per-prompt list of extra free-text documents
        # (e.g. a session document) — each {key, token, label, placeholder}.
        p.setdefault("extra_inputs", [])
        k = p["category"]
        if k not in cats:
            cats[k] = {"key": k, "title": p["category_title"],
                       "filter_label": _PROMPT_LIB_LABELS.get(k, p["category_title"]),
                       "prompts": []}
            order.append(k)
        cats[k]["prompts"].append(p)
    return prompts, [cats[k] for k in order]


PROMPT_LIBRARY, PROMPT_CATEGORIES = _load_prompt_library()
PROMPTS_BY_ID = {p["id"]: p for p in PROMPT_LIBRARY}

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_to_text(raw):
    """Pull readable text out of an uploaded .docx (paragraphs + table cells).
    The Claude API reads PDFs natively but not .docx, so Word resumes are
    flattened to text here before they go to the model."""
    from docx import Document
    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _build_library_messages(prompt, fields, resume, job_description, cover_letter, extras=None):
    """Fill the prompt template with the client's inputs and return an Anthropic
    messages list. Short [FIELD] placeholders and the large [PASTE …] blocks are
    substituted inline; a PDF resume rides along as a native document block."""
    template = prompt["template"]
    for key, val in (fields or {}).items():
        if val:
            template = template.replace(f"[{key}]", val)
    if job_description:
        template = template.replace("[PASTE THE JOB DESCRIPTION HERE]", job_description)
    if cover_letter:
        template = template.replace("[PASTE YOUR COVER LETTER HERE]", cover_letter)
    # Per-prompt extra free-text inputs (session document, interviewer, …).
    # Empty optional ones become "(not provided)" so no raw token reaches Claude.
    for ex in prompt.get("extra_inputs", []):
        val = ((extras or {}).get(ex["key"]) or "").strip()
        template = template.replace(ex["token"], val if val else "(not provided)")

    attachments, resume_text = [], ""
    if resume:
        if resume.get("kind") == "text":
            resume_text = (resume.get("text") or "").strip()
        elif resume.get("kind") == "file":
            data = resume.get("data") or ""
            if data.startswith("data:") and "," in data:
                data = data.split(",", 1)[1]  # strip data-URL prefix
            mime = resume.get("mime", "")
            name = (resume.get("name") or "").lower()
            if mime == "application/pdf" or name.endswith(".pdf"):
                attachments.append({"type": "document", "source": {
                    "type": "base64", "media_type": "application/pdf", "data": data}})
            else:
                raw = base64.b64decode(data)
                if mime == _DOCX_MIME or name.endswith(".docx"):
                    resume_text = _docx_to_text(raw)
                else:
                    resume_text = raw.decode("utf-8", "ignore").strip()

    if resume_text:
        if "[PASTE YOUR RESUME HERE]" in template:
            template = template.replace("[PASTE YOUR RESUME HERE]", resume_text)
        else:
            template += f"\n\nRESUME:\n\n{resume_text}"
    elif attachments and "[PASTE YOUR RESUME HERE]" in template:
        template = template.replace("[PASTE YOUR RESUME HERE]", "(resume attached as a PDF)")

    content = (attachments + [{"type": "text", "text": template}]) if attachments else template
    return [{"role": "user", "content": content}]


def _render_prompt_library():
    return render_template("prompt_library.html",
                           categories=PROMPT_CATEGORIES,
                           total=len(PROMPT_LIBRARY),
                           turnstile_site_key=TURNSTILE_SITE_KEY)


@app.route("/prompt-library")
def prompt_library():
    return _render_prompt_library()


# ── Generation rate limiting ───────────────────────────────────────────────
# The generate endpoint spends real money per call and is public, so cap usage
# per visitor and overall. State is in-memory and per-worker (gunicorn runs 2),
# so the effective global ceiling is ~2x PL_RATE_GLOBAL_DAY — fine as a first
# guardrail; move to a Redis-backed limiter if you need an exact shared ceiling.
# Bot protection (e.g. Cloudflare Turnstile) is the next layer and needs your
# Cloudflare keys to wire up. Tune all limits via env vars.
PL_RATE_PER_MIN = int(os.environ.get("PL_RATE_PER_MIN", "8"))
PL_RATE_PER_DAY = int(os.environ.get("PL_RATE_PER_DAY", "40"))
PL_RATE_GLOBAL_DAY = int(os.environ.get("PL_RATE_GLOBAL_DAY", "600"))
_pl_ip_hits = defaultdict(deque)
_pl_global = deque()
_pl_lock = threading.Lock()


def _client_ip():
    fwd = request.headers.get("X-Forwarded-For", "")
    return (fwd.split(",")[0].strip() if fwd else request.remote_addr) or "unknown"


def _rate_limit_message():
    """Record this request and return a user-facing message if the visitor (or
    the service overall) is over a limit, else None."""
    now = time.time()
    ip = _client_ip()
    with _pl_lock:
        dq = _pl_ip_hits[ip]
        while dq and now - dq[0] > 86400:
            dq.popleft()
        while _pl_global and now - _pl_global[0] > 86400:
            _pl_global.popleft()
        if sum(1 for t in dq if now - t <= 60) >= PL_RATE_PER_MIN:
            return "You're going a little fast — please wait a minute and try again."
        if len(dq) >= PL_RATE_PER_DAY:
            return "You've reached today's limit for AI generations. Please try again tomorrow."
        if len(_pl_global) >= PL_RATE_GLOBAL_DAY:
            return "This tool is at capacity right now — please try again in a little while."
        dq.append(now)
        _pl_global.append(now)
    return None


def _verify_turnstile(token):
    """True if the request may proceed. No secret configured -> always True.
    Secret set but token missing/invalid -> False. Network/verification errors
    fail OPEN (allow) — the per-IP rate limiter remains the cost backstop, so a
    brief Cloudflare outage won't block real clients."""
    if not TURNSTILE_SECRET:
        return True
    if not token:
        return False
    payload = urllib.parse.urlencode({
        "secret": TURNSTILE_SECRET, "response": token, "remoteip": _client_ip(),
    }).encode()
    try:
        with urllib.request.urlopen(_TURNSTILE_VERIFY_URL, data=payload, timeout=10) as resp:
            return bool(json.loads(resp.read().decode()).get("success"))
    except Exception:
        return True  # fail open on infra errors; rate limiter still applies


@app.route("/api/prompt-library/generate", methods=["POST"])
def prompt_library_generate():
    data = request.get_json(silent=True) or {}
    prompt = PROMPTS_BY_ID.get(data.get("id"))
    if not prompt:
        return jsonify({"error": "Unknown prompt."}), 404

    if not _verify_turnstile(data.get("turnstile_token")):
        return jsonify({"error": "Verification failed — please refresh the page and try again."}), 403

    limited = _rate_limit_message()
    if limited:
        return jsonify({"error": limited}), 429

    resume = data.get("resume") or {}
    if resume.get("kind") == "file" and len(resume.get("data") or "") > 14_000_000:
        return jsonify({"error": "That file is too large — please use one under 10 MB."}), 400

    try:
        messages = _build_library_messages(
            prompt, data.get("fields"), resume,
            (data.get("job_description") or "").strip(),
            (data.get("cover_letter") or "").strip(),
            data.get("extras"),
        )
    except Exception as e:
        return jsonify({"error": f"Could not read your upload: {e}"}), 400

    model = MODEL_SMART if prompt.get("model") == "smart" else MODEL_FAST
    max_tokens = 8000 if prompt.get("model") == "smart" else 4000

    def gen():
        try:
            with client.messages.stream(model=model, max_tokens=max_tokens, messages=messages) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps(text)}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"event: error\ndata: {json.dumps(str(e))}\n\n"

    return Response(stream_with_context(gen()), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── Session Generator ─────────────────────────────────────────────────────────

@app.route("/session-generator")
def session_generator():
    return render_template("session_generator.html")


@app.route("/api/session/generate", methods=["POST"])
def session_generate():
    data = request.get_json()
    prework = (data.get("prework") or "").strip()
    resume = (data.get("resume") or "").strip()

    if not prework or len(prework) < 100:
        return jsonify({"error": "Client pre-work is required."}), 400

    client_input = (
        f"PREWORK:\n\n{prework}\n\n---\n\nRESUME:\n\n{resume}"
        if resume
        else f"PREWORK:\n\n{prework}\n\n(No resume provided -- work from pre-work only)"
    )

    def _sse(event, data_obj):
        return f"event: {event}\ndata: {json.dumps(data_obj)}\n\n"

    def generate():
        yield _sse("step", {"step": 1})
        synthesis_chunks = []
        try:
            with client.messages.stream(
                model=MODEL_SMART,
                max_tokens=12000,
                system=SYNTHESIS_SYSTEM,
                messages=[{"role": "user", "content":
                    f"Here is the client data:\n\n{client_input}\n\nProduce the complete synthesis."}],
            ) as stream:
                for text in stream.text_stream:
                    synthesis_chunks.append(text)
                    yield _sse("ping", {})
        except Exception as e:
            yield _sse("error", {"message": str(e)})
            return

        synthesis = "".join(synthesis_chunks)
        yield _sse("step", {"step": 2})

        guide_chunks = []
        try:
            with client.messages.stream(
                model=MODEL_SMART,
                max_tokens=32000,
                system=COACH_GUIDE_SYSTEM,
                messages=[{"role": "user", "content":
                    f"Here is the client data:\n\n{client_input}\n\n---\n\nSYNTHESIS:\n\n{synthesis}\n\nProduce the complete coach session guide."}],
            ) as stream:
                for text in stream.text_stream:
                    guide_chunks.append(text)
                    yield _sse("guide_chunk", {"text": text})
        except Exception as e:
            yield _sse("error", {"message": str(e)})
            return

        guide = "".join(guide_chunks)
        yield _sse("step", {"step": 3})

        worksheet_chunks = []
        try:
            with client.messages.stream(
                model=MODEL_FAST,
                max_tokens=16000,
                system=WORKSHEET_SYSTEM,
                messages=[{"role": "user", "content":
                    f"Here is the client data:\n\n{client_input}\n\n---\n\nSYNTHESIS:\n\n{synthesis}\n\n---\n\nCOACH GUIDE:\n\n{guide}\n\nProduce the lean session worksheet."}],
            ) as stream:
                for text in stream.text_stream:
                    worksheet_chunks.append(text)
                    yield _sse("ping", {})
        except Exception as e:
            yield _sse("error", {"message": str(e)})
            return

        yield _sse("done", {"guide": guide, "worksheet": "".join(worksheet_chunks)})

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


def _session_slug(data):
    first = (data.get("first_name") or "client").strip().lower()
    last  = (data.get("last_name") or "").strip().lower()
    return re.sub(r"\s+", "_", f"{first}_{last}".strip("_")) or "client"


def _client_full_name(data):
    return " ".join(p for p in [(data.get("first_name") or "").strip(),
                                (data.get("last_name") or "").strip()] if p) or "Client"


@app.route("/api/session/reference.pdf", methods=["POST"])
def session_reference_pdf():
    data       = request.get_json()
    guide_md   = data.get("guide", "")
    name       = _client_full_name(data)
    slug       = _session_slug(data)
    pdf_bytes  = render_pdf(build_client_html(guide_md, name, "Coach Reference"))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_coach_reference.pdf")


@app.route("/api/session/worksheet.pdf", methods=["POST"])
def session_worksheet_pdf():
    data          = request.get_json()
    worksheet_md  = data.get("worksheet", "")
    name          = _client_full_name(data)
    slug          = _session_slug(data)
    pdf_bytes     = render_pdf(build_worksheet_html(worksheet_md, name))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_session_worksheet.pdf")


@app.route("/api/session/cue", methods=["POST"])
def session_cue_create():
    """Session Cue Screen — render the worksheet as a flip-through cue page, save
    it to the volume under an unguessable token with a short expiry, and return a
    shareable link (plus a QR code) to open on a tablet during the session."""
    data         = request.get_json()
    worksheet_md = data.get("worksheet", "")
    name         = _client_full_name(data)
    if not worksheet_md.strip():
        return jsonify({"error": "No worksheet to build a cue screen from."}), 400
    html = build_worksheet_cue_html(worksheet_md, name)
    token, expires = cue_store.save_cue(html, client_name=name)
    # Behind Railway's proxy request.scheme can be http; trust the forwarded proto
    # so the QR encodes https (phone cameras are finicky about http links).
    scheme = request.headers.get("X-Forwarded-Proto", request.scheme)
    url = f"{scheme}://{request.host}/cue/{token}"
    # Render a high-res PNG (crisp black/white squares, no SVG scaling quirks) and
    # hand it back as a data URI. error="l" keeps the module count low; border=4
    # is the required quiet zone.
    buf = io.BytesIO()
    segno.make(url, error="l").save(buf, kind="png", scale=10, border=4,
                                    dark="#1e2022", light="#ffffff")
    qr_data = "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode("ascii")
    return jsonify({"url": url, "expires": expires, "qr_data": qr_data})


def _cue_message_page(title, message):
    return Response(
        f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title><style>
body{{font-family:-apple-system,Segoe UI,Roboto,sans-serif;background:#f5f3f0;
  color:#2c3035;display:flex;align-items:center;justify-content:center;
  min-height:100vh;margin:0;padding:24px;}}
.box{{background:#fff;border:1px solid #e0dbd4;border-radius:8px;padding:40px 44px;
  max-width:440px;text-align:center;}}
h1{{font-size:20px;margin:0 0 10px;color:#1e2022;}}
p{{font-size:15px;line-height:1.6;color:#6a7078;margin:0;}}
</style></head><body><div class="box"><h1>{title}</h1><p>{message}</p></div></body></html>""",
        mimetype="text/html", status=410)


@app.route("/cue/<token>", methods=["GET"])
def session_cue_view(token):
    """Serve a saved cue page if the token is valid and not expired."""
    html = cue_store.load_cue(token)
    if html is None:
        return _cue_message_page(
            "Link expired or not found",
            "This session cue link has expired or doesn't exist. Generate a new "
            "one from the session screen.")
    return Response(html, mimetype="text/html")


@app.route("/api/session/guide.docx", methods=["POST"])
def session_guide_docx():
    data = request.get_json()
    guide_md = data.get("guide", "")
    slug = _session_slug(data)
    docx_bytes = markdown_to_docx(guide_md)
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{slug}_coach_guide.docx",
    )


# ── Session Timer ─────────────────────────────────────────────────────────────

@app.route("/session-timer")
def session_timer():
    return render_template("session_timer.html")

@app.route("/intake-prompter")
def intake_prompter():
    return render_template("intake_prompter.html")


# ── Session Builder ───────────────────────────────────────────────────────────

@app.route("/session-builder")
def session_builder():
    return render_template("session_builder.html")


@app.route("/api/builder/generate", methods=["POST"])
def builder_generate():
    data        = request.get_json()
    transcript  = (data.get("transcript") or "").strip()
    resume      = (data.get("resume") or "").strip()
    intake      = (data.get("intake") or "").strip()
    client_name = (data.get("client_name") or "").strip()
    client_id   = (data.get("client_id") or "").strip()

    if not transcript:
        return jsonify({"error": "Transcript is required."}), 400
    if not client_name:
        return jsonify({"error": "Client name is required."}), 400

    def _user_msg():
        msg = f"Here is the session transcript:\n\n{transcript}"
        if resume:
            msg += f"\n\n---\nRESUME:\n\n{resume}"
        if intake:
            msg += f"\n\n---\nINTAKE ANSWERS:\n\n{intake}"
        return msg

    def _sse(event, obj):
        return f"event: {event}\ndata: {json.dumps(obj)}\n\n"

    # Each job runs in its own thread and posts results to a queue
    q = queue.Queue()

    JOBS = [
        ("ip",          interview_program_prompt(client_name), 16000),
        ("stories",     stories_prompt(client_name),           16000),
        ("positioning", positioning_prompt(client_name, bool(resume), bool(intake)), 16000),
    ]

    def run_job(job_id, system_prompt, max_tokens):
        try:
            if job_id == "positioning":
                # Positioning returns a large structured JSON object. Route it
                # through complete_json so a truncated/malformed reply self-heals
                # and a genuine failure surfaces as an error instead of being
                # silently dropped at package-build time.
                data = complete_json(system_prompt, _user_msg(), MODEL_SMART, max_tokens)
                problems = positioning_warnings(data)
                if problems:
                    # Don't save thin positioning silently — surface it loudly so
                    # the coach re-runs instead of shipping a half-empty section.
                    raise ValueError("Positioning came back incomplete: " + " ".join(problems))
                content = json.dumps(data)
            else:
                content = complete_text(system_prompt, _user_msg(), MODEL_SMART, max_tokens)
            q.put(("done", job_id, content))
        except Exception as e:
            q.put(("error", job_id, str(e)))

    def run_resume_job():
        try:
            # Step 1: diagnostic brief
            brief_raw = complete_text(resume_diagnostic_prompt(), _user_msg(), MODEL_FAST, 1500)
            brief_raw = brief_raw.replace("```json", "").replace("```", "").strip()

            # Step 2: rewrite informed by brief
            rewrite_msg = f"STRATEGY BRIEF:\n{brief_raw}\n\n---\n\n{_user_msg()}"
            content = complete_text(resume_rewrite_prompt(), rewrite_msg, MODEL_SMART, 8000)
            q.put(("done", "resume", content))
        except Exception as e:
            q.put(("error", "resume", str(e)))

    def generate():
        # Signal all jobs as running
        for job_id, _, _ in JOBS:
            yield _sse("status", {"id": job_id, "state": "running"})
        yield _sse("status", {"id": "resume", "state": "running"})

        # Launch all jobs in parallel threads
        threads = []
        for job_id, system_prompt, max_tokens in JOBS:
            t = threading.Thread(target=run_job, args=(job_id, system_prompt, max_tokens), daemon=True)
            t.start()
            threads.append(t)
        t = threading.Thread(target=run_resume_job, daemon=True)
        t.start()
        threads.append(t)

        # Stream results as each job completes
        remaining = len(JOBS) + 1  # +1 for resume
        while remaining > 0:
            try:
                event_type, job_id, payload = q.get(timeout=300)
            except queue.Empty:
                yield _sse("error", {"id": "all", "message": "Timeout waiting for results."})
                return

            remaining -= 1
            if event_type == "done":
                try:
                    strategy_store.save_piece(client_name, job_id, payload, client_id=client_id)
                except Exception:
                    pass
                yield _sse("result", {"id": job_id, "state": "done", "content": payload})
            else:
                yield _sse("result", {"id": job_id, "state": "error", "message": payload})

        yield _sse("complete", {})

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/api/builder/resume.docx", methods=["POST"])
def builder_resume_docx():
    data        = request.get_json()
    resume_text = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    font        = (data.get("font") or "Calibri").strip()

    docx_bytes = resume_to_docx(resume_text, font=font)

    # Build filename: LastName_FirstName_Resume.docx
    parts = client_name.split()
    if len(parts) >= 2:
        filename = f"{parts[-1]}_{parts[0]}_Resume.docx"
    else:
        filename = f"{client_name.replace(' ', '_')}_Resume.docx"

    return send_file(
        io.BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/api/builder/ip.pdf", methods=["POST"])
def builder_ip_pdf():
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pdf_bytes   = render_pdf(build_client_html(content, client_name, "Interview Program"))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_interview_program.pdf")


@app.route("/api/builder/stories.pdf", methods=["POST"])
def builder_stories_pdf():
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pdf_bytes   = render_pdf(build_client_html(content, client_name, "Success Stories"))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_success_stories.pdf")


@app.route("/api/builder/positioning.pdf", methods=["POST"])
def builder_positioning_pdf():
    data        = request.get_json()
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pos_data    = data.get("positioning")
    if isinstance(pos_data, str):
        pos_data = json.loads(pos_data)
    pdf_bytes = render_pdf(build_positioning_html(pos_data, client_name))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_where_to_look.pdf")


# ── Session Builder Part 2 ────────────────────────────────────────────────────

@app.route("/session-builder-part2")
def session_builder_part2():
    return render_template("session_builder_part2.html")


@app.route("/api/builder2/generate", methods=["POST"])
def builder2_generate():
    data        = request.get_json()
    transcript  = (data.get("transcript") or "").strip()
    resume      = (data.get("resume") or "").strip()
    intake      = (data.get("intake") or "").strip()
    client_name = (data.get("client_name") or "").strip()
    client_id   = (data.get("client_id") or "").strip()

    if not transcript:
        return jsonify({"error": "Transcript is required."}), 400
    if not client_name:
        return jsonify({"error": "Client name is required."}), 400

    def _user_msg():
        msg = f"Here is the session transcript:\n\n{transcript}"
        if resume:
            msg += f"\n\n---\nRESUME:\n\n{resume}"
        if intake:
            msg += f"\n\n---\nINTAKE ANSWERS:\n\n{intake}"
        return msg

    def _sse(event, obj):
        return f"event: {event}\ndata: {json.dumps(obj)}\n\n"

    q = queue.Queue()

    def run_job(job_id, system_prompt, max_tokens):
        try:
            # Training is coach-only internal notes; everything else is client-facing.
            job_model = MODEL_FAST if job_id == "training" else MODEL_SMART
            content = complete_text(system_prompt, _user_msg(), job_model, max_tokens)
            q.put(("done", job_id, content))
        except Exception as e:
            q.put(("error", job_id, str(e)))

    def run_roleplay_job():
        try:
            part_a = complete_text(roleplay_a_prompt(client_name), _user_msg(), MODEL_SMART, 16000)
            part_b = complete_text(roleplay_b_prompt(client_name), _user_msg(), MODEL_SMART, 16000)
            q.put(("done", "roleplay", part_a + "\n\n" + part_b))
        except Exception as e:
            q.put(("error", "roleplay", str(e)))

    def run_resume_job():
        try:
            diag_msg = f"RESUME:\n\n{resume}" if resume else "No resume provided."
            if transcript:
                diag_msg += f"\n\n---\nSESSION TRANSCRIPT (Primary Truth — overrides resume where they conflict):\n\n{transcript}"
            diag = complete_json(DIAGNOSTIC_SYSTEM, diag_msg, MODEL_FAST, 2000)

            rw_msg = f"CANDIDATE LEVEL: {diag.get('level','EXECUTIVE')}\n"
            rw_msg += f"READINESS SCORE: {diag.get('readinessScore', 5)}/10\n"
            rw_msg += f"REVISION MODE: {diag.get('revisionMode','REWRITE')}\n\n"
            rw_msg += f"STRATEGY BRIEF:\n"
            rw_msg += f"Strengths to Preserve: {diag.get('strengths','')}\n"
            rw_msg += f"Narrative Arc: {diag.get('narrativeArc','')}\n"
            rw_msg += f"Impact Opportunities: {diag.get('impactOpportunities','')}\n"
            rw_msg += f"Keyword Alignment: {diag.get('keywordAlignment','')}\n"
            rw_msg += f"Positioning Hypothesis: {diag.get('positioningHypothesis','')}\n\n"
            if resume:
                rw_msg += f"RESUME:\n\n{resume}"
            if transcript:
                rw_msg += f"\n\n---\nSESSION TRANSCRIPT (Primary Truth):\n\n{transcript}"
            content = complete_text(REWRITE_SYSTEM, rw_msg, MODEL_SMART, 8000)
            q.put(("done", "resume", content))
        except Exception as e:
            q.put(("error", "resume", str(e)))

    JOBS = [
        ("summary",  summary_prompt(client_name),                          12000),
        ("branding", branding_prompt(client_name, bool(resume), bool(intake)), 12000),
        ("linkedin", linkedin_strategy_prompt(client_name, bool(resume), bool(intake)), 8000),
        ("training", training_prompt(client_name),                          8000),
        ("handoff",  coach_handoff_prompt(client_name),                     3000),
    ]

    def generate():
        for job_id, _, _ in JOBS:
            yield _sse("status", {"id": job_id, "state": "running"})
        yield _sse("status", {"id": "roleplay", "state": "running"})
        yield _sse("status", {"id": "resume", "state": "running"})

        threads = []
        for job_id, system_prompt, max_tokens in JOBS:
            t = threading.Thread(target=run_job, args=(job_id, system_prompt, max_tokens), daemon=True)
            t.start()
            threads.append(t)
        t = threading.Thread(target=run_roleplay_job, daemon=True)
        t.start()
        threads.append(t)
        t = threading.Thread(target=run_resume_job, daemon=True)
        t.start()
        threads.append(t)

        remaining = len(JOBS) + 2
        while remaining > 0:
            try:
                event_type, job_id, payload = q.get(timeout=600)
            except queue.Empty:
                yield _sse("error", {"id": "all", "message": "Timeout waiting for results."})
                return
            remaining -= 1
            if event_type == "done":
                try:
                    strategy_store.save_piece(client_name, job_id, payload, client_id=client_id)
                except Exception:
                    pass
                yield _sse("result", {"id": job_id, "state": "done", "content": payload})
            else:
                yield _sse("result", {"id": job_id, "state": "error", "message": payload})

        yield _sse("complete", {})

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/api/builder2/summary.pdf", methods=["POST"])
def builder2_summary_pdf():
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pdf_bytes   = render_pdf(build_client_html(content, client_name, "Strategy Session Summary"))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_session_summary.pdf")


@app.route("/api/builder2/roleplay.pdf", methods=["POST"])
def builder2_roleplay_pdf():
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pdf_bytes   = render_pdf(build_client_html(content, client_name, "Roleplay Interview Analysis"))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_roleplay_analysis.pdf")


@app.route("/api/builder2/branding.pdf", methods=["POST"])
def builder2_branding_pdf():
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pdf_bytes   = render_pdf(build_client_html(content, client_name, "Branding Profile"))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_branding_profile.pdf")


@app.route("/api/builder2/linkedin.pdf", methods=["POST"])
def builder2_linkedin_pdf():
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pdf_bytes   = render_pdf(build_client_html(content, client_name, "LinkedIn Strategy"))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_linkedin_strategy.pdf")


@app.route("/api/builder2/resume.docx", methods=["POST"])
def builder2_resume_docx():
    data        = request.get_json()
    content     = data.get("content", "")
    font        = (data.get("font") or "Calibri").strip()
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    docx_bytes  = resume_to_docx(content, font=font)
    return send_file(io.BytesIO(docx_bytes), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True, download_name=f"{slug}_resume.docx")


@app.route("/api/builder2/training.pdf", methods=["POST"])
def builder2_training_pdf():
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    pdf_bytes   = render_pdf(build_client_html(
        content, client_name, "Training Assessment — Coach Reference",
        eyebrow="Coach Reference",
    ))
    return send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_training_assessment.pdf")


@app.route("/api/builder2/training.html", methods=["POST"])
def builder2_training_html():
    import markdown as md_lib
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    from datetime import date
    today = date.today().strftime("%B %d, %Y")
    rendered = md_lib.markdown(content, extensions=["extra"])
    html_out = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Training Assessment — {client_name}</title>
<link href="https://fonts.googleapis.com/css2?family=Roboto+Slab:wght@400;700&family=Roboto:wght@300;400;600&family=Roboto+Mono&display=swap" rel="stylesheet">
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{font-family:'Roboto',sans-serif;font-size:15px;line-height:1.7;color:#2c3035;background:#f5f3f0;padding:40px 20px;}}
.doc{{max-width:780px;margin:0 auto;background:#fff;border:1px solid #ddd;border-radius:4px;overflow:hidden;}}
.doc-header{{background:#1e2022;padding:28px 40px 22px;}}
.doc-header-label{{font-family:'Roboto Mono',monospace;font-size:10px;letter-spacing:.18em;text-transform:uppercase;color:#5579a0;margin-bottom:6px;}}
.doc-header-name{{font-family:'Roboto Slab',serif;font-size:26px;font-weight:700;color:#fff;line-height:1.1;}}
.doc-header-sub{{font-size:12px;font-weight:300;color:#9ba0a6;margin-top:4px;}}
.doc-body{{padding:36px 40px 48px;}}
.doc-body h1{{font-family:'Roboto Slab',serif;font-size:20px;font-weight:700;color:#1e2022;margin:32px 0 10px;padding-bottom:7px;border-bottom:2px solid #3a5a7c;}}
.doc-body h1:first-child{{margin-top:0;}}
.doc-body h2{{font-family:'Roboto Slab',serif;font-size:16px;font-weight:600;color:#3a5a7c;margin:22px 0 7px;}}
.doc-body h3{{font-size:14px;font-weight:600;color:#1e2022;margin:16px 0 5px;}}
.doc-body p{{margin-bottom:11px;}}
.doc-body ul,.doc-body ol{{margin:6px 0 11px 22px;}}
.doc-body li{{margin-bottom:5px;}}
.doc-body blockquote{{border-left:3px solid #3a5a7c;padding:10px 0 10px 18px;margin:14px 0;background:#eef4fb;border-radius:0 2px 2px 0;}}
.doc-body blockquote p{{margin-bottom:0;font-style:italic;}}
.doc-body strong{{font-weight:600;color:#1e2022;}}
.doc-body hr{{border:none;border-top:1px solid #e0dbd4;margin:22px 0;}}
.doc-footer{{padding:14px 40px;border-top:1px solid #e0dbd4;font-family:'Roboto Mono',monospace;font-size:10px;color:#9ba0a6;letter-spacing:.08em;display:flex;justify-content:space-between;}}
</style></head><body>
<div class="doc">
  <div class="doc-header">
    <div class="doc-header-label">Coach Reference Only</div>
    <div class="doc-header-name">{client_name}</div>
    <div class="doc-header-sub">Training Assessment</div>
  </div>
  <div class="doc-body">{rendered}</div>
  <div class="doc-footer"><span>{client_name} — Training Assessment</span><span>{today}</span></div>
</div>
</body></html>"""
    return send_file(
        io.BytesIO(html_out.encode("utf-8")),
        mimetype="text/html",
        as_attachment=True,
        download_name=f"{slug}_training_assessment.html",
    )


@app.route("/api/builder2/handoff.html", methods=["POST"])
def builder2_handoff_html():
    import markdown as md_lib
    from datetime import date
    data        = request.get_json()
    content     = data.get("content", "")
    client_name = (data.get("client_name") or "client").strip()
    slug        = re.sub(r"[^a-z0-9_]", "", client_name.lower().replace(" ", "_"))
    today       = date.today().strftime("%B %d, %Y")
    rendered    = md_lib.markdown(content, extensions=["extra", "nl2br"])
    html_out = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Implementation Handoff — {client_name}</title>
<link href="https://fonts.googleapis.com/css2?family=Roboto+Slab:wght@400;700&family=Roboto:wght@300;400;600&family=Roboto+Mono&display=swap" rel="stylesheet">
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{font-family:'Roboto',sans-serif;font-size:15px;line-height:1.7;color:#2c3035;background:#f5f3f0;padding:40px 20px;}}
.doc{{max-width:720px;margin:0 auto;background:#fff;border:1px solid #ddd;border-radius:4px;overflow:hidden;}}
.doc-header{{background:#1e2022;padding:28px 40px 22px;}}
.doc-header-label{{font-family:'Roboto Mono',monospace;font-size:10px;letter-spacing:.18em;text-transform:uppercase;color:#5579a0;margin-bottom:6px;}}
.doc-header-name{{font-family:'Roboto Slab',serif;font-size:24px;font-weight:700;color:#fff;line-height:1.1;}}
.doc-header-sub{{font-size:12px;font-weight:300;color:#9ba0a6;margin-top:4px;}}
.doc-body{{padding:34px 40px 46px;}}
.doc-body p{{margin-bottom:13px;}}
.doc-body strong{{font-weight:600;color:#1e2022;}}
.doc-footer{{padding:14px 40px;border-top:1px solid #e0dbd4;font-family:'Roboto Mono',monospace;font-size:10px;color:#9ba0a6;letter-spacing:.08em;display:flex;justify-content:space-between;}}
</style></head><body>
<div class="doc">
  <div class="doc-header">
    <div class="doc-header-label">Coach Reference Only</div>
    <div class="doc-header-name">{client_name}</div>
    <div class="doc-header-sub">Implementation Coach Handoff</div>
  </div>
  <div class="doc-body">{rendered}</div>
  <div class="doc-footer"><span>{client_name} — Implementation Handoff</span><span>{today}</span></div>
</div>
</body></html>"""
    return send_file(
        io.BytesIO(html_out.encode("utf-8")),
        mimetype="text/html",
        as_attachment=True,
        download_name=f"{slug}_implementation_handoff.html",
    )


# ── All-in-one (full day: both runs + package build) ────────────────────────────

@app.route("/full-session")
def full_session_page():
    return render_template("full_session.html")


@app.route("/api/full/generate", methods=["POST"])
def full_generate():
    data        = request.get_json()
    transcript  = (data.get("transcript") or "").strip()
    resume      = (data.get("resume") or "").strip()
    intake      = (data.get("intake") or "").strip()
    client_name = (data.get("client_name") or "").strip()
    client_id   = (data.get("client_id") or "").strip()

    if not transcript:
        return jsonify({"error": "Transcript is required."}), 400
    if not client_name:
        return jsonify({"error": "Client name is required."}), 400

    def _user_msg():
        msg = f"Here is the session transcript:\n\n{transcript}"
        if resume:
            msg += f"\n\n---\nRESUME:\n\n{resume}"
        if intake:
            msg += f"\n\n---\nINTAKE ANSWERS:\n\n{intake}"
        return msg

    def _sse(event, obj):
        return f"event: {event}\ndata: {json.dumps(obj)}\n\n"

    q = queue.Queue()

    # Single-call jobs: (id, system_prompt, max_tokens, model)
    JOBS = [
        ("ip",          interview_program_prompt(client_name),                       16000, MODEL_SMART),
        ("stories",     stories_prompt(client_name),                                 16000, MODEL_SMART),
        ("summary",     summary_prompt(client_name),                                 12000, MODEL_SMART),
        ("branding",    branding_prompt(client_name, bool(resume), bool(intake)),    12000, MODEL_SMART),
        ("positioning", positioning_prompt(client_name, bool(resume), bool(intake)), 16000, MODEL_SMART),
        ("linkedin",    linkedin_strategy_prompt(client_name, bool(resume), bool(intake)), 8000, MODEL_SMART),
        ("training",    training_prompt(client_name),                                 8000, MODEL_FAST),
        ("handoff",     coach_handoff_prompt(client_name),                            3000, MODEL_SMART),
    ]

    def run_job(job_id, system_prompt, max_tokens, model):
        try:
            if job_id == "positioning":
                # Positioning returns a large structured JSON object. Route it
                # through complete_json so a truncated/malformed reply self-heals
                # and a genuine failure surfaces as an error instead of being
                # silently dropped at package-build time.
                data = complete_json(system_prompt, _user_msg(), model, max_tokens)
                problems = positioning_warnings(data)
                if problems:
                    # Don't save thin positioning silently — surface it loudly so
                    # the coach re-runs instead of shipping a half-empty section.
                    raise ValueError("Positioning came back incomplete: " + " ".join(problems))
                q.put(("done", job_id, json.dumps(data)))
                return
            content = complete_text(system_prompt, _user_msg(), model, max_tokens)
            q.put(("done", job_id, content))
        except Exception as e:
            q.put(("error", job_id, str(e)))

    def run_roleplay_job():
        try:
            part_a = complete_text(roleplay_a_prompt(client_name), _user_msg(), MODEL_SMART, 16000)
            part_b = complete_text(roleplay_b_prompt(client_name), _user_msg(), MODEL_SMART, 16000)
            q.put(("done", "roleplay", part_a + "\n\n" + part_b))
        except Exception as e:
            q.put(("error", "roleplay", str(e)))

    def run_resume_job():
        try:
            brief = complete_text(resume_diagnostic_prompt(), _user_msg(), MODEL_FAST, 1500)
            brief = brief.replace("```json", "").replace("```", "").strip()
            rewrite_msg = f"STRATEGY BRIEF:\n{brief}\n\n---\n\n{_user_msg()}"
            content = complete_text(resume_rewrite_prompt(), rewrite_msg, MODEL_SMART, 8000)
            q.put(("done", "resume", content))
        except Exception as e:
            q.put(("error", "resume", str(e)))

    def generate():
        all_ids = [j[0] for j in JOBS] + ["roleplay", "resume"]
        for jid in all_ids:
            yield _sse("status", {"id": jid, "state": "running"})

        threads = []
        for job_id, system_prompt, max_tokens, model in JOBS:
            t = threading.Thread(target=run_job, args=(job_id, system_prompt, max_tokens, model), daemon=True)
            t.start(); threads.append(t)
        for runner in (run_roleplay_job, run_resume_job):
            t = threading.Thread(target=runner, daemon=True)
            t.start(); threads.append(t)

        remaining = len(JOBS) + 2
        while remaining > 0:
            try:
                event_type, job_id, payload = q.get(timeout=600)
            except queue.Empty:
                yield _sse("error", {"id": "all", "message": "Timeout waiting for results."})
                return
            remaining -= 1
            if event_type == "done":
                try:
                    strategy_store.save_piece(client_name, job_id, payload, client_id=client_id)
                except Exception:
                    pass
                yield _sse("result", {"id": job_id, "state": "done", "content": payload})
            else:
                yield _sse("result", {"id": job_id, "state": "error", "message": payload})

        yield _sse("complete", {})

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Strategy Package (combined client book) ─────────────────────────────────────

@app.route("/strategy-package")
def strategy_package_page():
    return render_template("strategy_package.html")


@app.route("/api/package/clients")
def package_clients():
    return jsonify(strategy_store.list_clients())


@app.route("/api/package/build.pdf", methods=["POST"])
def package_build_pdf():
    data        = request.get_json()
    client_name = (data.get("client_name") or "").strip()
    # The package page sends back the storage slug from list_clients() so we
    # load the exact folder that was saved (identifier slug or legacy name
    # slug); fall back to the name slug for older clients that send only a name.
    storage_key = (data.get("client_key") or "").strip() or strategy_store.slugify(client_name)
    if not client_name:
        return jsonify({"error": "Client name is required."}), 400

    pieces = strategy_store.load_all(storage_key)
    if not pieces:
        return jsonify({"error": f"No saved content found for {client_name}. Run the Session Builder for this client first."}), 404

    html      = build_strategy_package_html(client_name, pieces)
    pdf_bytes = render_pdf(html)
    slug      = storage_key
    resp = send_file(io.BytesIO(pdf_bytes), mimetype="application/pdf",
                     as_attachment=True, download_name=f"{slug}_strategy_package.pdf")
    # Pre-flight warnings (thin/invalid positioning, etc.) ride along in a
    # header so the UI can flag problems before this package reaches a client.
    warnings = package_warnings(pieces)
    if warnings:
        resp.headers["X-Package-Warnings"] = json.dumps(warnings)
        resp.headers["Access-Control-Expose-Headers"] = "X-Package-Warnings"
    return resp


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=os.environ.get("FLASK_DEBUG", "false").lower() == "true")
